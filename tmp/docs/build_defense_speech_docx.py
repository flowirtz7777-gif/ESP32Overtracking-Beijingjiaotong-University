from pathlib import Path
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(os.environ["OUT_DOCX_TMP"])


def set_east_asia_font(run, font_name="Microsoft YaHei"):
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)


def set_style_font(style, font_name="Microsoft YaHei", size=Pt(11), color=None, bold=None):
    style.font.name = font_name
    style.font.size = size
    if color:
        style.font.color.rgb = RGBColor(*color)
    if bold is not None:
        style.font.bold = bold
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)


def add_page_turn(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("（翻页）")
    set_east_asia_font(r)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(192, 0, 0)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.style = "Heading 2"
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_east_asia_font(r)


def add_body(doc, text):
    p = doc.add_paragraph()
    p.style = "Normal"
    p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_east_asia_font(r)


SCRIPT = [
    (
        "封面页",
        [
            "各位老师好，我是关霁恩。我们汇报的项目是“基于多传感融合与短时运动学预测的大型、铰接车辆右转内轮差盲区实时预警系统”。下面我将从研究背景、现有方案不足、项目目标、技术路线、创新点以及实施计划六个方面进行汇报。"
        ],
    ),
    (
        "研究背景",
        [
            "首先是研究背景。本项目面向大型车辆，尤其是半挂车辆在低速右转过程中的内轮差盲区风险。大型车辆右转时，前后轮轨迹并不重合，车辆右侧会形成内轮差危险区域；而半挂车还存在牵引车和挂车姿态耦合，危险区域会随着转角、铰接角和速度实时变化。",
            "对于行人、电动车等弱势交通参与者来说，一旦进入车辆右侧盲区，驾驶员往往难以及时感知。因此，我们需要的不是事后解释，而是提前给出可解释的空间风险判断。",
        ],
    ),
    (
        "既有方案与关键不足",
        [
            "现有方案主要包括单一雷达或视觉预警、静态盲区提示、单车体简化模型以及轨迹控制类模型。它们各自有作用，但也存在明显断层：有的只能感知近旁目标，难以表达车辆未来扫掠区域；有的采用固定安全带，忽略了车辆转角、速度和铰接角的实时变化；还有一些模型更偏向轨迹生成或控制输入，并不等同于面向行人的短时风险判定。"
        ],
    ),
    (
        "问题重新定义",
        [
            "因此，我们将问题重新定义为一个短时空间风险预测问题。也就是说，本项目不是判断驾驶员长期意图，也不是控制车辆如何转弯，而是在车辆当前状态已知的情况下，预测未来短时间内右侧车身会扫过哪里，并判断目标是否会进入危险区域。核心表达就是：未来短时内，目标是否进入车辆右侧扫掠危险区。"
        ],
    ),
    (
        "项目目标与研究内容",
        [
            "围绕这个问题，我们的研究目标是构建一套面向半挂货车低速右转过程的实时预警原型系统。系统采集车速、前轮转角、铰接角以及右侧目标信息，建立牵引车和挂车耦合运动学模型，预测未来短时右侧扫掠危险区，并形成 0 到 3 级分级报警。",
            "对应的关键问题包括：半挂车右转内轮差风险如何建模，如何建立考虑铰接约束的短时运动学预测模型，如何构造多时间尺度危险区，以及如何在 ESP32-S3 等低成本平台上实现实时判警。",
        ],
    ),
    (
        "总体技术路线",
        [
            "我们的总体技术路线可以概括为六步。第一，输入实时状态，包括车速 v、前轮转角 alpha 和铰接角 phi；第二，通过耦合运动学递推车辆未来短时状态；第三，计算牵引车前后轴中点、牵引销点和挂车后轴中点等关键几何点；第四，生成三层右侧扫掠危险区；第五，将雷达和视觉目标统一到车体坐标系下进行风险评估；第六，输出 0 到 3 级声光报警。"
        ],
    ),
    (
        "危险区构建与分级报警",
        [
            "在危险区构建上，我们把内轮差风险转化为未来短时右侧车身扫掠区域。系统生成 0.3 秒、1.0 秒和 2.0 秒三层窗口，分别对应立即危险区、报警区和预警区。若目标命中当前车身占用区或 0.3 秒窗口，则判为最高风险；若进入 1.0 秒或 2.0 秒窗口，则分别给出报警和预警。",
            "同时，我们加入报警状态机滞回设计，例如升级需要连续满足条件，降级需要连续不满足条件，从而降低传感器噪声造成的报警抖动。",
        ],
    ),
    (
        "系统架构与验证工具链",
        [
            "在系统实现上，我们采用低成本双 MCU 架构。ESP32-S3 作为主控平台，负责传感器采集、短时预测、风险评估和报警输出；OpenMV H7 Plus 作为视觉协处理器，用于右侧目标识别。输入部分包括前轮转角、铰接角、CAN 或 OBD 车速，以及 HLK-LD2450 毫米波雷达目标检测。",
            "目前我们已经搭建了仿真与验证工具链，包括 PID 工况仿真导出器、MATLAB 回放、TTC 预警复盘、转弯全过程盲区分析软件和 CFG 配置生成器。这里需要说明，PID 数据不是现实驾驶真值，而是用于生成平滑、连续、可控、可复现的虚拟传感器输入，帮助验证主预测模型在多场景下的稳定性和有效性。",
        ],
    ),
    (
        "创新点与项目特色",
        [
            "本项目的创新点主要体现在四个方面。第一，把内轮差风险从固定盲区或单一距离，转化为面向实时预警的时空扫掠危险区；第二，建立考虑铰接结构约束的半挂车短时耦合运动学模型；第三，基于 0.3 秒、1.0 秒和 2.0 秒时间窗口形成多层危险区和分级报警机制；第四，形成一条仿真驱动、可解释、可移植到低成本嵌入式平台的原型路线。",
        ],
    ),
    (
        "实施计划与预期成果",
        [
            "后续我们将按计划推进模型验证、硬件标定、嵌入式算法移植、雷达视觉融合联调和半实物测试。预期成果包括一套半挂货车右转内轮差盲区实时预警原型系统，一套可复现的短时耦合运动学预测算法，三层右侧扫掠危险区构建方法，以及配套的 MATLAB 仿真、网页分析工具、技术说明和演示材料。",
            "以上就是我们的汇报，恳请各位老师批评指正。",
        ],
    ),
]


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    set_style_font(doc.styles["Normal"], "Microsoft YaHei", Pt(11), color=(0, 0, 0))
    set_style_font(doc.styles["Heading 1"], "Microsoft YaHei", Pt(16), color=(46, 116, 181), bold=True)
    set_style_font(doc.styles["Heading 2"], "Microsoft YaHei", Pt(13), color=(46, 116, 181), bold=True)

    header_p = section.header.paragraphs[0]
    header_p.text = ""
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = header_p.add_run("开题答辩演讲稿")
    set_east_asia_font(hr)
    hr.font.size = Pt(9)
    hr.font.color.rgb = RGBColor(100, 100, 100)

    footer_p = section.footer.paragraphs[0]
    footer_p.text = ""
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run("北京交通大学本科生创新训练项目")
    set_east_asia_font(fr)
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(100, 100, 100)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("答辩演讲稿（3-4分钟版）")
    set_east_asia_font(r)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(31, 77, 120)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("基于多传感融合与短时运动学预测的大型（铰接）车辆右转内轮差盲区实时预警系统")
    set_east_asia_font(r)
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(80, 80, 80)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("使用说明：正文中的“（翻页）”为现场换页提示，不需要读出。整体语速按正常答辩节奏约 3-4 分钟。")
    set_east_asia_font(r)
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(90, 90, 90)

    for idx, (heading, paras) in enumerate(SCRIPT):
        add_section_heading(doc, heading)
        for para in paras:
            add_body(doc, para)
        if idx != len(SCRIPT) - 1:
            add_page_turn(doc)

    doc.save(OUT)


if __name__ == "__main__":
    build_docx()
    print(OUT)
