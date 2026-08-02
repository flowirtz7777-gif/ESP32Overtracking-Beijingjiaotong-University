from pathlib import Path
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


DOCX_PATH = Path(os.environ["DOCX_PATH"])


def set_east_asia_font(run, font_name="Microsoft YaHei"):
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    set_east_asia_font(r)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(31, 77, 120)


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    set_east_asia_font(r)
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(90, 90, 90)


def add_qa(doc, number, question, answer):
    q = doc.add_paragraph()
    q.paragraph_format.space_before = Pt(7)
    q.paragraph_format.space_after = Pt(2)
    qr = q.add_run(f"{number}. {question}")
    set_east_asia_font(qr)
    qr.bold = True
    qr.font.size = Pt(11)
    qr.font.color.rgb = RGBColor(31, 77, 120)

    a = doc.add_paragraph()
    a.paragraph_format.left_indent = Pt(18)
    a.paragraph_format.first_line_indent = Pt(0)
    a.paragraph_format.line_spacing = 1.12
    a.paragraph_format.space_after = Pt(3)
    ar = a.add_run("参考回答：" + answer)
    set_east_asia_font(ar)
    ar.font.size = Pt(10.5)


QA_ITEMS = [
    (
        "你们这个项目和普通盲区报警有什么区别？",
        "普通盲区报警更多是判断“当前旁边有没有目标”，而我们关注的是“未来短时间内车辆右侧会扫过哪里”。也就是把内轮差风险从静态盲区，转化为可预测的短时扫掠危险区。",
    ),
    (
        "为什么说这不是一个控制问题？",
        "因为我们不控制车辆怎么转弯，也不规划驾驶轨迹。我们只根据当前车速、前轮转角和铰接角，预测未来右侧车身扫掠区域，用于风险判断和报警。",
    ),
    (
        "为什么要考虑半挂车，而不是普通货车？",
        "半挂车存在牵引车和挂车之间的铰接关系，挂车后部轨迹更复杂，内轮差风险也更明显。如果只用普通单车体模型，会忽略挂车姿态变化，风险判断不够准确。",
    ),
    (
        "你们的核心输入有哪些？",
        "主要输入包括车速 v、前轮转角 alpha、铰接角 phi，以及右侧雷达和视觉检测到的目标位置。前三个用于预测车辆未来姿态，目标信息用于风险判定。",
    ),
    (
        "为什么需要铰接角 phi？",
        "因为半挂车的挂车姿态不完全由牵引车航向决定，铰接角会直接影响挂车后部的扫掠轨迹。没有 phi，就很难准确判断挂车右侧是否会扫到目标。",
    ),
    (
        "三层危险区为什么设为 0.3 秒、1.0 秒、2.0 秒？",
        "这三个时间窗分别对应紧急、报警和预警三个层级。0.3 秒表示几乎立即发生风险，1 秒用于强报警，2 秒用于提前提醒，便于形成分级响应。",
    ),
    (
        "如何判断目标的风险等级？",
        "如果目标在当前车身占用区或 0.3 秒危险区内，判为最高风险；如果进入 1.0 秒或 2.0 秒扫掠区，则分别判为报警或预警。同时结合 TTC 估计，提高对动态目标的判断能力。",
    ),
    (
        "为什么还要加入报警状态机？",
        "传感器会有噪声，如果风险等级一变就立刻报警，容易抖动。状态机通过连续满足才升级、连续不满足才降级，让报警更稳定。",
    ),
    (
        "你们为什么不用纯视觉或者纯雷达？",
        "单一传感器都有局限。视觉有类别识别优势，但受光照影响；毫米波雷达对距离和速度更稳定，但目标语义弱。融合后可以提升可靠性。",
    ),
    (
        "OpenMV 的作用是什么？ESP32 不能直接跑视觉吗？",
        "ESP32-S3 主要负责实时采集、预测和报警，不适合承担较重的视觉识别任务。OpenMV 作为视觉协处理器，可以把识别结果传给 ESP32，降低主控负担。",
    ),
    (
        "你们的模型和机器学习方法相比有什么优势？",
        "我们采用运动学模型，优点是可解释、计算量小、适合嵌入式部署。机器学习可能需要大量真实场景数据，而本项目目前更适合先做低成本、可验证的原型系统。",
    ),
    (
        "PID 在项目中到底起什么作用？",
        "PID 不是最终在线判警模型，也不是现实驾驶真值。它只是用来生成平滑、连续、可控的仿真输入，帮助验证主预测模型。",
    ),
    (
        "如果 PID 数据不是真实驾驶数据，验证还有意义吗？",
        "有意义，但要明确边界。PID 数据可以用于早期算法闭环验证，检查模型递推、危险区生成和报警逻辑是否正确；后续还需要半实物和真实车辆测试进一步验证。",
    ),
    (
        "如何证明你们的预测结果是准确的？",
        "前期可以通过 MATLAB 仿真、PID 工况回放、TTC 复盘和盲区分析工具进行一致性验证。后续会通过传感器标定、半实物实验和实车低速场景测试进一步验证。",
    ),
    (
        "这个系统实时性够吗？",
        "够。核心计算是运动学递推、多边形构建和点在多边形内判断，计算量较小，适合 ESP32-S3 这类低成本 MCU 实现。",
    ),
    (
        "如果传感器误差比较大怎么办？",
        "首先要做传感器标定和滤波；其次报警逻辑中会加入滞回机制，减少瞬时误差带来的误报警。后续也可以通过安全裕度或阈值调节提高鲁棒性。",
    ),
    (
        "你们项目最大的技术难点是什么？",
        "难点主要有三个：半挂车铰接运动学建模、右侧扫掠危险区的实时构建，以及雷达和视觉目标统一到车体坐标系后的稳定分级判警。",
    ),
    (
        "创新点到底在哪里？",
        "主要是把内轮差从固定盲区或单一距离问题，转化为未来短时右侧车身扫掠危险区预测问题；同时结合半挂车铰接模型、三层时间窗和低成本嵌入式实现路线。",
    ),
    (
        "你们目前完成到了什么程度？",
        "目前已经完成项目问题定义、核心算法路线、仿真验证工具链和原型系统架构设计。后续重点是硬件标定、ESP32 端算法移植、雷达视觉融合和半实物测试。",
    ),
    (
        "如果最后实车测试效果不理想，怎么改进？",
        "我们会先区分问题来源：是传感器误差、坐标标定误差、模型参数误差，还是报警阈值设置不合理。然后分别通过标定、滤波、参数修正和阈值优化迭代，而不是直接推翻整体路线。",
    ),
]


def main():
    doc = Document(DOCX_PATH)

    doc.add_page_break()
    add_title(doc, "答辩问答准备（20问）")
    add_note(doc, "使用说明：本部分为答辩备答材料，不需要在正式汇报中朗读。回答时可先给结论，再补充一两句理由。")

    for idx, (question, answer) in enumerate(QA_ITEMS, start=1):
        add_qa(doc, idx, question, answer)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
