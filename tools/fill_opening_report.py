from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]


def find_template() -> Path:
    for path in ROOT.rglob("*.docx"):
        if path.name.startswith("~$") or "node_modules" in str(path):
            continue
        if path.stat().st_size == 67841:
            return path
    raise FileNotFoundError("Could not locate 大创项目立项报告_模板.docx")


def set_font(run, size=10.5, bold=False):
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph(para, text, size=12, bold=False, align=None):
    para.text = text
    if align is not None:
        para.alignment = align
    for run in para.runs:
        set_font(run, size=size, bold=bold)


def set_cell(cell, text, font_size=10.5, bold=False, align=None):
    cell.text = ""
    parts = text.split("\n") if text else [""]
    for index, part in enumerate(parts):
        para = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.15
        if align is not None:
            para.alignment = align
        run = para.add_run(part)
        set_font(run, size=font_size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def main():
    template = find_template()
    out_path = template.parent / "开题报告_半挂货车右转内轮差预警系统_填写稿.docx"
    doc = Document(template)

    for style_name in ["Normal"]:
        style = doc.styles[style_name]
        style.font.name = "SimSun"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(10.5)

    project_title = "基于多传感融合与短时运动学预测的半挂货车右转内轮差盲区实时预警系统"

    cover = {
        6: (project_title, 18, True, WD_ALIGN_PARAGRAPH.CENTER),
        9: ("项目管理单位：交通运输学院（待填写）", 12, False, None),
        10: ("项目类型：☑ 创新训练项目", 12, False, None),
        11: ("□ 创业训练项目", 12, False, None),
        12: ("□ 创业实践项目", 12, False, None),
        13: ("负责学生：待填写（学号：待填写）", 12, False, None),
        14: ("项目成员：待填写、待填写", 12, False, None),
        15: ("指导教师：待填写", 12, False, None),
        19: ("2026年6月", 12, False, WD_ALIGN_PARAGRAPH.CENTER),
    }
    for idx, args in cover.items():
        set_paragraph(doc.paragraphs[idx], *args)

    # 一、基本情况
    t0 = doc.tables[0]
    set_cell(t0.cell(0, 2), project_title, bold=True)
    set_cell(
        t0.cell(1, 2),
        "☑实物作品  □发表论文  ☑软件程序  ☑研究报告\n☑设计方案  □影音作品  □其他",
        font_size=10,
    )
    set_cell(t0.cell(2, 2), "一年")
    member_rows = [
        ("1主持", "待填写", "待填写", "智能运输工程 / 自动化 / 电子信息（待填写）", "交通运输学院（待填写）", "待填写"),
        ("2参加", "待填写", "待填写", "待填写", "待填写", "待填写"),
        ("3参加", "待填写", "待填写", "待填写", "待填写", "待填写"),
    ]
    for row_index, values in enumerate(member_rows, start=4):
        for col_index, value in enumerate(values, start=1):
            set_cell(t0.cell(row_index, col_index), value, font_size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    teacher_values = ["1", "待填写", "待填写", "车辆工程 / 交通安全 / 嵌入式系统（待填写）", "待填写", "待填写"]
    for col_index, value in enumerate(teacher_values, start=1):
        set_cell(t0.cell(8, col_index), value, font_size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)

    # 二、项目简介
    intro = (
        "随着城市物流、港口运输和干线货运车辆规模持续扩大，重型货车特别是半挂车在交叉口右转时形成的内轮差盲区，"
        "已成为行人、非机动车与大型车辆混行场景中的典型安全风险。传统后视镜、盲区镜和单一静态报警装置主要依赖驾驶员观察或固定区域触发，"
        "难以表达半挂车右侧车身边界在未来短时间内的动态扫掠范围，容易出现提前量不足或误报警问题。\n"
        "本项目面向半挂货车低速右转过程中的“卷入—碾压”风险，拟构建一套基于车辆短时运动学预测、毫米波雷达与视觉感知融合的实时预警原型系统。"
        "系统以ESP32-S3为主控，采集车速、前轮转向角、铰接角及右侧目标信息，建立牵引车—挂车耦合运动学模型，"
        "预测未来0.3 s、1.0 s、2.0 s三层右侧车身扫掠危险区，并对行人、电动车等目标进行分级判警。"
        "项目预期形成可演示的软硬件样机、MATLAB/网页仿真验证工具链和研究报告，为大型车辆右转主动安全预警提供低成本、可解释、可部署的技术方案。"
    )
    set_cell(doc.tables[1].cell(0, 0), intro)

    # 三、立项依据
    basis = """（一）行业研究现状和发展动态
重型货车右转内轮差风险具有低速、近距离、遮挡强和反应时间短等特点。现有安全辅助方案主要包括盲区镜、摄像头显示、超声波/毫米波雷达探测、基于固定盲区的声光报警等。这类方案能够提升驾驶员感知能力，但多数仍停留在“当前目标是否位于固定危险区域”的层面，对半挂车铰接角变化、挂车后部内切轨迹及未来短时扫掠区域表达不足。国内外关于车辆转弯运动学、铰接车内轮差区域和扫掠包络已有一定理论基础，但面向低成本嵌入式实时预警的工程化方案仍需解决三个问题：一是如何用车载实时可测量的车速、转向角和铰接角构造可解释预测模型；二是如何将“内轮差”从单一距离指标转化为未来短时车身扫掠危险区；三是如何在算力有限的单片机上实现多目标分级判警和稳定报警。
本项目已具备一定前期基础：已形成半挂车耦合运动学建模思路，明确了牵引车后轴中心、牵引销、挂车后轴等关键点位关系；已完成MATLAB阶段的运动学步进、三层扫掠多边形生成、点在多边形内判定和PID工况导出器等仿真基础；已规划ESP32-S3、OpenMV、AS5600、铰接角电位器、CAN车速和HLK-LD2450雷达等硬件架构。尚需补充的是硬件联调、传感器标定、目标融合、真实/半实物场景验证和报警策略参数整定。

（二）主要研究内容
1. 半挂货车右转内轮差风险机理分析：梳理牵引车与挂车低速转弯的几何关系，明确车体坐标系、车速、前轮转向角、铰接角、牵引销位置和挂车后轴位置等变量定义。
2. 短时耦合运动学预测模型：建立包含牵引车后轴至牵引销距离项的铰接车运动学步进模型，预测未来短时间内牵引车与挂车姿态演化。
3. 右侧扫掠危险区构建：基于A/B/H/T关键点和车身宽度，分别构建0.3 s、1.0 s、2.0 s三层危险区，形成紧急碰撞区、报警区和预警区。
4. 多传感器目标感知与融合：利用毫米波雷达获取右侧目标距离与速度，利用OpenMV视觉模块识别行人、非机动车等目标，并统一映射到车体坐标系。
5. 风险评估与报警状态机：通过点在扫掠区内判定和目标恒速外推TTC估计，输出0-3级风险，并设计带滞回的声光报警逻辑。
6. 仿真、半实物与样机验证：构建PID工况生成、MATLAB回放、网页盲区分析和ESP32端实时运行的验证链路，比较主预测模型与静态危险带、简化单车模型等基线方案。

（三）创新点与项目特色
1. 将内轮差风险表达为“未来短时右侧车身扫掠危险区”，而不是单一几何距离或固定盲区带，能够更直接服务实时预警。
2. 面向半挂车铰接结构建立耦合运动学预测模型，保留后轴至牵引销距离对挂车角速度的影响，避免将半挂车简单等效为单体车辆。
3. 采用三层时间窗分级判警，将几何扫掠区、目标外推和报警状态机结合，兼顾提前量、误报警控制和紧急响应。
4. 采用ESP32-S3 + OpenMV + 毫米波雷达的低成本双MCU架构，计算链路可解释、可移植，适合竞赛样机和后续工程转化。
5. 构建“PID工况生成—主预测模型—扫掠区判定—仿真/网页复盘”的闭环验证工具链，在缺少真实车辆条件下也能开展可复现的模型验证。"""
    set_cell(doc.tables[2].cell(0, 0), basis, font_size=10)

    # 四、项目实施方案
    plan = """（一）研究方案
1. 理论基础与总体思路
本项目将右转内轮差预警定义为短时风险预测问题，而不是车辆控制问题。系统不尝试预测驾驶员完整意图或规划长期轨迹，而是以当前可测的车速v、前轮转向角alpha、铰接角phi为输入，通过低速刚体运动学递推得到未来短时间内牵引车与挂车姿态，再由右侧车身边界生成危险区，与雷达/视觉目标进行空间和时间上的重叠判定。

2. 关键模型与算法
（1）半挂车耦合运动学模型。以牵引车后轴中心B为车体坐标原点，状态量为[xB, yB, theta, phi]。牵引车横摆角速度由车速与前轮转角确定；挂车角速度由牵引车后轴速度、铰接角以及牵引销偏置共同确定，递推得到牵引车航向、铰接角和挂车航向。
（2）扫掠危险区生成。由牵引车前轴点A、后轴点B、牵引销点H、挂车后轴点T推导车身右侧边界。A/B/H使用牵引车航向投影右侧边界，T使用挂车航向投影右侧边界。按预测时间窗构建PolyI、PolyA、PolyW三层扫掠区域。
（3）目标风险判定。将毫米波雷达与视觉识别目标统一到车体坐标系，对目标当前位置和短时恒速外推位置进行点在多边形内判定，并结合TTC阈值输出风险等级。报警状态机采用连续升级、延迟降级的滞回策略，降低瞬时抖动造成的误报警。

3. 硬件与软件实现
硬件侧采用ESP32-S3作为主控，负责传感器采集、运动学预测、风险判定和声光报警；OpenMV H7 Plus作为视觉协处理器，负责行人/非机动车目标检测并通过UART上传；AS5600采集前轮转向角，电位器采集铰接角，CAN/OBD采集车速，HLK-LD2450毫米波雷达采集右侧目标位置和速度。软件侧分为传感器驱动、滤波与标定、预测器、风险评估、报警状态机、日志与调试接口等模块。

4. 验证方法
前期利用PID工况仿真导出器生成平滑的v、alpha、phi序列，将其作为虚拟传感器输入喂给主预测模型，验证运动学递推和扫掠区生成的几何一致性；中期利用MATLAB与网页盲区分析工具构造边界目标点，统计不同预测窗口和状态机策略下的提前量；后期搭建桌面/小车半实物测试环境，验证雷达目标触发、视觉目标识别、报警等级切换和端到端实时性。

5. 人员分工与资源条件
项目负责人负责总体方案、运动学模型、进度协调和报告撰写；成员1负责MATLAB/网页仿真、数据记录和对照实验；成员2负责ESP32固件、传感器驱动、报警输出和硬件联调。指导教师负责研究路线把关、实验安全审核和阶段成果评审。计划使用学院实验室的基础电子制作工具、计算机、开源软件环境、低速测试场地以及学校图书馆/数据库资源。"""
    set_cell(doc.tables[3].cell(0, 0), plan, font_size=10)

    progress = [
        ("1", "需求分析与资料调研：完成内轮差事故场景、铰接车运动学、雷达/视觉预警方案调研，明确项目指标和技术路线。", "2026.06-2026.07", "形成调研记录和开题材料"),
        ("2", "模型建立与仿真验证：完善半挂车耦合运动学模型、三层扫掠区生成、点在多边形内判定和PID工况回放链路。", "2026.07-2026.08", "输出MATLAB仿真结果"),
        ("3", "硬件选型与传感器标定：完成ESP32-S3、OpenMV、AS5600、铰接角传感器、CAN车速、毫米波雷达的接口测试。", "2026.08-2026.09", "完成硬件接线和标定方案"),
        ("4", "嵌入式算法移植：将预测器、风险评估、报警状态机移植到ESP32端，完成50 Hz主循环和串口日志。", "2026.09-2026.10", "形成可运行固件"),
        ("5", "雷达视觉融合与报警联调：实现目标坐标统一、类别/位置融合、风险等级输出和LED/蜂鸣器报警策略。", "2026.10-2026.11", "完成桌面联调演示"),
        ("6", "半实物场景测试：构建右转边界目标点和典型行人/电动车目标场景，测试提前量、误报警和漏报警情况。", "2026.11-2026.12", "形成测试数据和问题清单"),
        ("7", "优化与成果整理：优化预测窗口、状态机参数和硬件封装，整理软件、样机、研究报告、展示材料。", "2026.12-2027.01", "完成中期/结题材料"),
        ("8", "结题验收与推广论证：完成最终演示、文档归档、专利/软著/论文可行性评估和后续转化计划。", "2027.01-2027.05", "提交结题报告和成果包"),
    ]
    for row_index, values in enumerate(progress, start=3):
        for col_index, value in enumerate(values):
            set_cell(doc.tables[3].cell(row_index, col_index), value, font_size=9)

    budget = """（三）经费预算及使用计划
项目经费主要用于样机硬件、实验耗材、资料打印与成果整理，预算可根据学院立项额度进一步调整。
1. 传感器与主控硬件：ESP32-S3开发板、OpenMV H7 Plus、HLK-LD2450毫米波雷达、AS5600角度传感器、角度电位器、CAN/TWAI接口模块、蜂鸣器、LED、电源模块、线束与连接器等，预计1800元。
2. 结构与实验耗材：传感器支架、外壳、安装板、3D打印/亚克力加工、低速测试标靶、反光/警示材料等，预计900元。
3. 软件测试与数据采集：存储卡、串口调试器、测试线材、日志记录介质、小型工具耗材等，预计500元。
4. 资料、打印和成果申报：调研资料打印、报告装订、展板制作、专利/软著前期检索与申请材料准备等，预计800元。
合计预计4000元。经费使用遵循实物优先、测试必需、凭据报销原则，优先保障核心传感器、嵌入式样机和验证场景搭建。"""
    set_cell(doc.tables[3].cell(12, 0), budget, font_size=10)

    # 五、项目预期成果
    expected = """项目实施后预期达到以下成果：
1. 完成一套半挂货车右转内轮差盲区实时预警原型系统，包括ESP32-S3主控、OpenMV视觉协处理、毫米波雷达、转向角/铰接角/车速采集和声光报警模块。
2. 形成可复现的短时耦合运动学预测算法，能够根据实时v、alpha、phi输入生成0.3 s、1.0 s、2.0 s三层右侧扫掠危险区，并输出0-3级目标风险。
3. 完成MATLAB仿真、PID工况回放、网页盲区分析和嵌入式端运行验证，形成若干典型右转场景的测试数据、图表和分析结论。
4. 输出项目研究报告、技术说明书、源代码和演示视频；具备申请软件著作权、实用新型专利或学生科技竞赛作品的基础。
5. 在样机演示层面实现对右侧行人/非机动车目标的分级预警，目标主循环周期约20 ms，核心算法适配低成本嵌入式平台。
6. 形成后续真实车辆低速测试、产品化封装和校企合作验证的技术储备。"""
    set_cell(doc.tables[4].cell(0, 0), expected, font_size=10)

    commitment = """团队成员承诺：
我们认真阅读了2026年大学生创新创业训练项目申报通知，符合相关申报条件，了解项目立项内容，立项报告书及大创项目管理系统中本团队成员和指导教师信息均将据实填写。
团队成员承诺对填写的各项内容真实性负责，项目立项报告内容原创，项目内容保证没有知识产权争议。如获准立项，我们承诺遵守学校相关规定，按计划认真开展研究和实践工作，按期完成模型验证、样机开发、测试记录和成果整理，取得预期成果。

项目负责人（签名）：

项目成员（签名）："""
    set_cell(doc.tables[5].cell(0, 0), commitment, font_size=10.5)

    teacher = """指导教师意见：
本项目面向重型/半挂货车右转内轮差盲区安全这一具有现实意义的交通安全问题，研究目标明确，技术路线较完整，能够综合训练学生在车辆运动学建模、嵌入式系统开发、多传感器感知、仿真验证和工程样机调试等方面的能力。项目拟采用短时运动学预测构建右侧车身扫掠危险区，并结合雷达、视觉目标信息进行分级预警，具有一定创新性和实践价值。
建议项目团队在实施过程中重点把握三点：一是严格区分仿真工况与真实车辆数据，确保实验结论表述客观；二是做好传感器标定、坐标系统一和报警阈值整定，保证样机演示稳定可靠；三是按阶段保存代码、数据和测试记录，形成可复现的研究成果。本人同意该项目申请立项。

指导教师（签名）：
                                                   年   月    日"""
    set_cell(doc.tables[6].cell(0, 0), teacher, font_size=10.5)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        set_font(run, size=run.font.size.pt if run.font.size else 10.5, bold=run.bold or False)

    doc.save(out_path)
    print(out_path)


if __name__ == "__main__":
    main()
